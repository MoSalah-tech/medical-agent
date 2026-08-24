import os
import logging
from typing import Optional

from langchain_core.documents import Document
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_postgres.vectorstores import PGVector
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
from docx import Document as DocxDocument
from PIL import Image
import pytesseract

from app.medical_agent.core.config import (
    gemini_api_key,
    gemini_embedding_model,
    database_url_psycopg,
    pgvector_collection,
)

logger = logging.getLogger(__name__)

_embeddings = None
_vectorstore = None

_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=150,
    separators=["\n\n", "\n", ". ", " ", ""],
)

class RAGError(Exception):
    pass


def get_embeddings() -> GoogleGenerativeAIEmbeddings:
    global _embeddings
    if _embeddings is None:
        _embeddings = GoogleGenerativeAIEmbeddings(
            model=gemini_embedding_model,
            google_api_key=gemini_api_key,
        )
    return _embeddings


def get_vectorstore() -> PGVector:
    """
    Lazily build and reuse a PGVector client.
    """
    global _vectorstore
    if _vectorstore is None:
        _vectorstore = PGVector(
            embeddings=get_embeddings(),
            collection_name=pgvector_collection,
            connection=database_url_psycopg,
            use_jsonb=True,
        )
    return _vectorstore


# ------------------------ Text extraction ---------------------------

def extract_pdf_text(pdf_path: str) -> str:
    reader = PdfReader(pdf_path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_docx_text(docx_path: str) -> str:
    doc = DocxDocument(docx_path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def extract_txt_text(txt_path: str) -> str:
    with open(txt_path, "r", encoding="utf-8") as f:
        return f.read()


def extract_image_text(image_path: str) -> str:
    if os.getenv("TESSERACT_CMD"):
        pytesseract.pytesseract.tesseract_cmd = os.getenv("TESSERACT_CMD")
    image = Image.open(image_path)
    return pytesseract.image_to_string(image)


def extract_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_pdf_text(file_path)
    elif ext == ".docx":
        return extract_docx_text(file_path)
    elif ext in [".txt", ".text"]:
        return extract_txt_text(file_path)
    elif ext in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
        return extract_image_text(file_path)
    else:
        raise RAGError(f"Unsupported file type: {ext}")


# ------------------------ Chunking & Ingestion ----------------------

def chunk_text(text: str, source: str, user_id: str) -> list[Document]:
    chunks = _splitter.split_text(text)
    return [
        Document(
            page_content=chunk,
            metadata={"source": source, "user_id": user_id, "chunk_index": i},
        )
        for i, chunk in enumerate(chunks)
    ]


def ingest_file(file_path: str, user_id: str, source_name: Optional[str] = None) -> int:
    """
    Extract text from a supported file, chunk it, embed chunks, and store
    them in pgvector. Returns the number of chunks stored.
    """
    text = extract_text(file_path)
    if not text.strip():
        logger.warning("No extractable text in %s", file_path)
        return 0

    docs = chunk_text(text, source=source_name or file_path, user_id=user_id)
    try:
        get_vectorstore().add_documents(docs)
    except Exception as exc:
        raise RAGError(f"Failed to store embeddings for {file_path}: {exc}") from exc
    return len(docs)


# Backward-compatible alias
def ingest_pdf(pdf_path: str, user_id: str, source_name: Optional[str] = None) -> int:
    return ingest_file(pdf_path, user_id, source_name)


# ------------------------ Retrieval --------------------------------

def search(query: str, user_id: str, top_k: int = 5) -> list[dict]:
    """
    Retrieve top-k chunks belonging to a specific user.
    """
    k = top_k if top_k is not None else 5
    results = get_vectorstore().similarity_search_with_score(
        query, k=k, filter={"user_id": user_id}
    )
    return [
        {
            "content": doc.page_content,
            "source": doc.metadata.get("source", "unknown"),
            "score": float(score),
            "metadata": doc.metadata,
        }
        for doc, score in results
    ]